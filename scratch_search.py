import os
import docx
import json

def search_in_docx(folder_path, keyword):
    results = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                path = os.path.join(root, file)
                try:
                    doc = docx.Document(path)
                    for i, p in enumerate(doc.paragraphs):
                        if keyword in p.text:
                            results.append(f"{file} (Paragraph {i}): {p.text.strip()}")
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if keyword in cell.text:
                                    results.append(f"{file} (Table cell): {cell.text.strip()}")
                except Exception as e:
                    pass
    return results

print("=== DOCX Search ===")
docx_results = search_in_docx('./docs', '중고차')
for r in set(docx_results):
    print(r)
