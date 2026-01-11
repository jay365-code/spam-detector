from docx import Document
import os

def extract_text_from_docx(docx_path, output_txt_path):
    if not os.path.exists(docx_path):
        print(f"Error: Docx file not found at {docx_path}")
        return

    print(f"Reading from {docx_path}...")
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract from tables if necessary, but start with paragraphs
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             full_text.append(cell.text)

    text_content = "\n".join(full_text)
    
    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write(text_content)
    
    print(f"Successfully converted to {output_txt_path}")
    print(f"Preview (first 200 chars):\n{text_content[:200]}")

if __name__ == "__main__":
    DOCX_FILE = r"../docs/스팸 및 햄 분류 가이드(20200112).docx"
    TXT_FILE = r"data/spam_guide.txt"
    extract_text_from_docx(DOCX_FILE, TXT_FILE)
